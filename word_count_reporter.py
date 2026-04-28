"""Word count reporter for text and DOCX documents.

This script reads an input file containing a project title and a list of
chapters with their corresponding file paths, counts the words in each
document (.txt or .docx), and generates an HTML report with a table of
word counts.

Usage:
    python word_count_reporter.py INPUT_FILE [OPTIONS]

Input file format:
    The input file must be parsable by the custom `inputfile` module.
    Expected structure: title line followed by chapter entries, each with
    a chapter name and file path.

Dependencies:
    - python-docx: for reading .docx files
    - beautifulsoup4: for HTML generation
    - inputfile (local module): for parsing the input file format
    - report_template.html: must exist in the working directory

Output:
    Generates an HTML report and opens it in the default web browser.
    Optionally backs up source files as .txt in a subdirectory.
"""

from datetime import datetime as dt
from docx import Document
from bs4 import BeautifulSoup
from typing import Any, Optional, Union
import webbrowser
import sys
import argparse
import inputfile
import shutil
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SCRIPT_DIR = Path(__file__).resolve().parent
REPORT_DIR = Path(
    "C:\\Users\\Boris\\Documents\\programming\\git repos\\word-count-reporter\\reports"
)
# dir where js and css to embed live
ASSETS_SRC = SCRIPT_DIR / "assets"

ENC = "utf-8"


def main(args: list[str]) -> None:
    """Main entry point: parse arguments, process files, generate report.

    Args:
        args (list): Command-line arguments (typically sys.argv[1:]).

    Returns:
        None
    """
    parser = argparse.ArgumentParser(
        description="Create a word count report",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument("input", type=Path, help="Input file (positional argument)")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        help="""Output file. If not supplied,
                        will be based on title and timestamp""",
        required=False,
    )
    parser.add_argument(
        "-b",
        "--backup",
        required=False,
        action="store_true",
        help="Backup input files as text files in the same "
        "directory holding the report. Notes: "
        "(1) if any of the files are .docx files, will "
        "convert them to text files; .txt files will be "
        "backed up as is; other file type not supported."
        " (2) if --backup and --output both given, "
        "then --output should be the directory for both "
        "the report and the backed up "
        "files, NOT the path to the report itself.",
    )
    parser.add_argument(
        "-t",
        "--notimestamp",
        required=False,
        action="store_true",
        help="Don't timestamp output file",
    )
    parser.add_argument(
        "-u",
        "--usetitle",
        required=False,
        action="store_true",
        help="When --output not given, use the project's "
        "title in filename generated for the report.",
    )
    parser.add_argument(
        "--loglevel", default="info", choices=["debug", "info"], help="Log level."
    )
    parser.add_argument(
        "-F",
        "--FORCE",
        required=False,
        action="store_true",
        default=False,
        help="Overwrite output file if exists",
    )
    args = parser.parse_args(args)

    # a logger for my debugging purposes
    loglevel = logging.DEBUG
    if args.loglevel:
        match args.loglevel:
            case "debug":
                loglevel = logging.DEBUG
            case "info":
                loglevel = logging.INFO
    logger.setLevel(loglevel)
    console = logging.StreamHandler()
    # console.setFormatter(logging.Formatter('%(name)-12s: %(message)s'))
    logger.addHandler(console)

    ifile = args.input
    if not ifile.exists():
        raise Exception(f"Input file {str(ifile)} does not exist!")
    if not ifile.is_absolute():
        ifile = (SCRIPT_DIR / ifile).resolve()

    # parse the input file
    # note: input_data is array of arrays.
    # each inner array corresponds to a chapter.
    # each inner array is [chap name, path to file]
    title, input_data = parse_input_file(ifile)

    # generic output file and possibly folder
    # (if --backup, will encorporate a folder)

    generic_filename = "word-count-report"
    generic_folder = "word_count_w_backup"
    ts = timestamp()
    if args.usetitle:
        safe_title = title.replace(" ", "_")
        generic_filename = safe_title + "-" + generic_filename
        generic_folder = safe_title + "-" + generic_folder
    if not args.notimestamp:
        generic_filename += "_" + ts
        generic_folder += "_" + ts
    generic_filename += ".html"

    # determine report file path, and backup dir if needed

    report_dir = None  # Path object to hold the report (and backups)
    report_filename = None  # name of the report

    if args.output:  # --output arg given
        output_path = args.output
        ext = output_path.suffix
        if args.backup:  # --backup given: --output should be a folder
            logger.debug("op #1: (--output, --backup)")
            if ext:
                raise Exception(
                    "\n--output and --backup given, but "
                    " --output appears to be a file (it has "
                    "an extension). If you're giving --backup "
                    "and --output, then --output should specify "
                    "a FOLDER, which will hold both the report "
                    "file as well as the backed up files :)"
                )
            report_dir = output_path
            report_filename = generic_filename
        else:  # --backup not given; --output should be a file (the report)
            logger.debug("op #2: (--output, no --backup)")
            if not ext:
                raise Exception("\n\n--output does not appear to be a file")
            report_dir = output_path.parent
            report_filename = output_path.name
    else:  # --output arg not given - use default
        report_dir = REPORT_DIR / title
        report_filename = generic_filename
        if (
            args.backup
        ):  # if --backup, default report dir should be an extra FOLDER: it will hold backup + report
            logger.debug("op #3: (no --output, --backup)")
            report_dir = report_dir / generic_folder

    # final path to the report
    report_file = report_dir / report_filename
    if not report_file.is_absolute():
        report_file = SCRIPT_DIR / report_file

    logger.debug(f"report dir     : {str(report_dir)}")
    logger.debug(f"report filename: {report_filename}")
    logger.debug(f"Report file    : {str(report_file)}")

    # if the report dir doesnt exist, create it
    report_dir.mkdir(parents=True, exist_ok=True)

    # 2:40 pm 10/10/24
    # 2:48

    logger.debug("original input data:")
    logger.debug(input_data)

    if args.backup:
        backup_files = backup(input_data, report_dir)
        new_idata = []
        logger.debug("loop through and back up files")
        # zip allows you to loop through 2 lists at once
        for original, backed_up in zip(input_data, backup_files):
            new_idata.append([original[0], backed_up])
        input_data = new_idata

    logger.debug("new input data:")
    logger.debug(input_data)

    report = str(make_report(title, input_data, report_file, args.FORCE))

    logger.info(report)
    webbrowser.open(report)


def parse_input_file(filepath: Path) -> tuple[str, list[list[Any]]]:
    """Parse the input file and extract title and chapter file paths.

    Args:
        filepath (Path): Path to the input file.

    Returns:
        tuple: (title, my_data) where:
            - title (str): Project title from input file.
            - my_data (list): List of [chapter_name (str), file_path (Path)] pairs.
    """
    title, file_data, inputfile_had_parts = inputfile.parse_data_file(str(filepath))
    my_data = []
    for data in file_data:
        filename = data[1]  # chapter name
        filepath = Path(data[2])  # filepath to chapter
        if not filename:  # if no chapter name, make it the name of the file
            filename = filepath.name
        my_data.append([filename, filepath])
    return title, my_data


def make_report(
    title: str, file_info_list: list[list[Any]], outfile: Path, force: bool
) -> Path:
    """Augment file info with word counts and generate the report.

    Args:
        title (str): Project title.
        file_info_list (list): List of [chapter_name (str), file_path (Path)] pairs.
        outfile (Path): Output HTML file path.
        force (bool): Overwrite output file if exists.

    Returns:
        Path: Path to the generated report file (returned by generate_report).
    """
    for file_info in file_info_list:
        filepath = Path(file_info[1])
        file_info.append(file_word_count(filepath))
    return generate_report(title, file_info_list, outfile, force)


def timestamp() -> str:
    """Return a timestamp string safe for use in filenames.

    Returns:
        str: Current timestamp in format YYYY_MM_DD-HH_MM_SS.
    """
    # format the datetime without any spaces
    fmt = "%Y_%m_%d-%H_%M_%S"
    ct = dt.now().strftime(fmt)
    return str(ct)


def date_string() -> str:
    """human readable/useful date string"""
    date = dt.now()
    formatted_date = date.strftime("%B %d, %Y")

    time = dt.now()
    formatted_time = time.strftime("%I:%M:%S %p")

    return f"{formatted_date}, {formatted_time}"


def file_contents(filepath: Path) -> str:
    """Read and return the entire contents of a text file.

    Args:
        filepath (Path): Path to the file to read.

    Returns:
        str: File contents.

    Raises:
        Exception: If the file does not exist.
    """
    if not filepath.exists():
        raise Exception(f"{str(filepath)} doesn't exist")
    file = open(filepath, "r", encoding=ENC)
    file_str = file.read()
    file.close()
    return file_str


def write_file(filepath: Path, data: str, force: bool) -> None:
    """Write data to a file, creating directories as needed.

    Args:
        filepath (Path): Absolute path to the output file.
        data (str): Content to write.
        force (bool): If True, overwrite existing file. If False, require
                      that the file does not already exist.

    Raises:
        Exception: If filepath is not absolute.
        Exception: If file exists and force is False.
    """
    if not filepath.is_absolute():
        raise Exception(f"Output file {str(filepath)} is not absolute!")
    if filepath.exists() and not force:
        raise Exception(
            f"Output file {str(filepath)} exists! " "Use --FORCE / -F" " to overwrite"
        )
    # create dirs in path if they don't exist
    basedir = filepath.parent
    basedir.mkdir(parents=True, exist_ok=True)

    mode = "w" if force else "x"
    with open(filepath, mode, encoding=ENC) as wr:
        wr.write(data)


def generate_report(
    title: str, word_count_info: list[Any], outfile: Path, force: bool
) -> Path:
    """Generate an HTML report from word count data.

    Args:
        title (str): Project title.
        word_count_info (list): List of [chapter_name (str), file_path (Path), word_count (int)].
        outfile (Path): Output HTML file path.
        force (bool): Overwrite output file if exists.

    Returns:
        Path: Path to the generated report file.
    """
    date = date_string()

    contents = file_contents(Path("report_template.html"))
    # css and js to embed
    css_contents = file_contents(ASSETS_SRC / "style.css")
    js_contents = file_contents(ASSETS_SRC / "scripts.js")

    contents = contents.replace("%title%", title)
    contents = contents.replace("%date%", date)
    # embed css and js for portability
    contents = contents.replace("%css%", css_contents)
    contents = contents.replace("%script%", js_contents)

    soup = BeautifulSoup(contents.encode(ENC), "html.parser")

    # Get <tbody> in word-count-table
    table = soup.find(attrs={"id": "word-count-table"})
    if not table:
        raise RuntimeError("Could not find table with id='word-count-table'")
    tbody = table.find("tbody")
    if not tbody:
        raise RuntimeError("Table missing required <tbody> element")

    # append rows directly to <tbody>
    total_words = 0
    for idx, count_info in enumerate(word_count_info):
        tbody.append(word_count_row(count_info, idx + 1))
        total_words += count_info[2]
    tbody.append(total_row(total_words))

    pretty_soup = soup.prettify(formatter="html")
    write_file(outfile, pretty_soup, force)
    return outfile


def file_url(filepath: Path) -> str:
    """Convert a local file path to a file:// URL for HTML links.

    Args:
        filepath (Path): Local file path (backslashes allowed).

    Returns:
        str: URL in format file:///C:/path/to/file.
    """
    filestr = str(filepath).replace("\\", "/")
    filestr = "file:///" + filestr
    return filestr


def number(numstr: Union[int, str]) -> str:
    """Format a number with thousand separators.

    Args:
        numstr (int or str): Number to format.

    Returns:
        str: Number with commas as thousand separators.
    """
    return f"{numstr:,}"


def word_count_row(file_info: list[Any], row_num: int) -> BeautifulSoup:
    """Create an HTML table row for a single chapter's word count.

    Args:
        file_info (list): [chapter_name (str), file_path (Path), word_count (int)].
        row_num (int): Row number (1-indexed) for display.

    Returns:
        BeautifulSoup: HTML element representing the table row.
    """
    chapter_name = file_info[0]
    chapter_path = file_url(file_info[1])
    word_count = number(file_info[2])

    return BeautifulSoup(
        f"""
        <tr>
            <td>{row_num}</td>
            <td>{chapter_name}</td>
            <td><a href="{chapter_path}" target=_blank>file</a></td>
            <td>{word_count} words</td>
        </tr>
        """,
        "html.parser",
    )


def total_row(total: int) -> BeautifulSoup:
    """Create an HTML table row for the total word count.

    Args:
        total (int): Total word count across all chapters.

    Returns:
        BeautifulSoup: HTML element representing the total row.
    """
    return BeautifulSoup(
        f"""
        <tr id="total">
            <td colspan="3">Total</td>
            <td>{number(total)}</td>
        </tr>
        """,
        "html.parser",
    )


def word_count_docx(filepath: Path) -> int:
    """Count words in a .docx file.

    Args:
        filepath (Path): Path to the .docx file.

    Returns:
        int: Number of words (split on whitespace per paragraph).
    """
    document = Document(str(filepath))
    num_words = 0
    for paragraph in document.paragraphs:
        words = paragraph.text.split()
        num_words += len(words)
    return num_words


def word_count_txt(filepath: Path) -> int:
    """Count words in a .txt file.

    Args:
        filepath (Path): Path to the .txt file.

    Returns:
        int: Number of words (split on whitespace per line).
    """
    num_words = 0
    with open(filepath, "r", encoding=ENC) as f:
        for line in f:
            words = line.split()
            num_words += len(words)
    return num_words


def file_word_count(filepath: Path) -> int:
    """Dispatch word counting based on file extension.

    Args:
        filepath (Path): Path to the file (.txt or .docx).

    Returns:
        int: Word count.

    Raises:
        Exception: If file does not exist or extension is not supported.
    """
    if not filepath.exists():
        raise Exception(f"File {str(filepath)} does not exist!")

    extension = filepath.suffix
    if extension == ".txt":
        return word_count_txt(filepath)
    if extension == ".docx":
        return word_count_docx(filepath)
    else:
        raise Exception(
            f"invalid filetype {extension}. "
            "Only .txt and .docx "
            "currently supported"
        )


def backup(files_info: list[list[Any]], report_dir: Path) -> list[Path]:
    """Back up all source files to a subdirectory.

    Args:
        files_info (list): List of [chapter_name (str), file_path (Path)] pairs.
        report_dir (Path): Directory where the report will be saved; backups
                          are placed in a "files" subdirectory.

    Returns:
        list[Path]: Paths to the backed-up files, in the same order as files_info.
    """
    destfiles = []
    backup_dir = report_dir / "files"
    for file_info in files_info:
        # 'backup' the file and add filepath of
        # backed up file to destfiles
        destfiles.append(backup_file(Path(file_info[1]), file_info[0], backup_dir))
    return destfiles


def backup_file(chapter_filepath: Path, chapter_name: str, backup_dir: Path) -> Path:
    """Back up a single file, converting .docx to .txt if needed.

    Args:
        chapter_filepath (Path): Path to the source file.
        chapter_name (str): Name of the chapter (used for naming .txt output
                            when converting from .docx).
        backup_dir (Path): Destination directory for the backup.

    Returns:
        Path: Path to the backed-up file.

    Raises:
        Exception: If file extension is not .txt or .docx.
    """
    logger.debug(
        f"\n\n*chapter name: {chapter_name}\n* srcfile: {str(chapter_filepath)}\n* dest dir: {str(backup_dir)}"
    )
    extension = chapter_filepath.suffix
    dest_filepath = None  # will be a Path object
    if extension == ".docx":
        filename_no_ext = Path(
            chapter_name
        ).stem  # explore: isnt' chapter_name just a simple string (not a path)?
        dest_filepath = backup_dir / Path(filename_no_ext + ".txt")
        docx_to_txt(chapter_filepath, dest_filepath)
    elif extension == ".txt":
        filename = chapter_filepath.name
        dest_filepath = backup_dir / filename
        shutil.copyfile(str(chapter_filepath), str(dest_filepath))
    else:
        raise Exception("file to backup isn't .docx or .txt")
    logger.debug(f"\tFile backed up to: {str(dest_filepath)}")
    return dest_filepath


def docx_to_txt(srcpath: Path, destpath: Path) -> None:
    """Convert a .docx file to a plain text file.

    Args:
        srcpath (Path): Path to the source .docx file.
        destpath (Path): Path for the output .txt file.

    Raises:
        Exception: If srcpath does not exist or destpath already exists.
    """
    if not srcpath.exists():
        raise Exception(
            f"Trying to convert a docx to txt, but the docx file doesn't exist: {str(srcpath)}"
        )
    if destpath.exists():
        raise Exception(
            f"Trying to convert a docx to a text file, "
            f"but proposed destination path already exists.\n\n"
            f"srcfile: {str(srcpath)}\ndest path: {str(destpath)}"
        )

    # need to create parent dir of file writing to or python will error
    destpath.parent.mkdir(parents=True, exist_ok=True)
    with open(destpath, "w", encoding=ENC) as destfile:
        document = Document(str(srcpath))
        for paragraph in document.paragraphs:
            destfile.write(paragraph.text + "\n")


if __name__ == "__main__":
    main(sys.argv[1:])
